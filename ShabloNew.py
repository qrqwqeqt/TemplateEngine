import re
from docx import Document

doc = Document(r"C:\Users\user\HowPrograming\Python\Shablonizator\text.docx")
textdocx = []
myDict = {}
varList = []

for paragraph in doc.paragraphs:
    words = re.split(r'(#\S+#)', paragraph.text)
    textdocx.append(words)
# print(words)


for number in range(len(words)):
    if words[number][0] == '#' and words[number][-1] == '#' and len(words[number]) > 2:
        if words[number] not in myDict:
            myDict[words[number]] = [number]
        else:
            myDict[words[number]].append(number)

# for key, value in myDict.items():
#     print(f'{key}: {value}')

def showKeys():
    for key in myDict:
        print(key)

while True:
    showKeys()
    req = str(input())
    if req in myDict:
        newInfo = input('Enter updated data in {} paragraph: '.format(req))
        for i in myDict[req]:
            words[i] = newInfo
        if input('Continue editing? (Yes/No) ') == 'No':
            break
    else:
        print('Key not found')
        showKeys()

print(words)

newDoc = Document()
text = ''.join(words)
newDoc.add_paragraph(text=text)
newDoc.save("text_1.docx")